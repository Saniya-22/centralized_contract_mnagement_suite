#!/usr/bin/env python3
"""
Generate docs/GOVGIG_Architecture_And_Env.docx from docs/GOVGIG_ARCHITECTURE.md.
Requires: pip install -r scripts/requirements-docs.txt
Run from repo root: python scripts/build_architecture_doc.py
"""

from pathlib import Path
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    print(
        "Missing python-docx. Run: pip install -r scripts/requirements-docs.txt",
        file=sys.stderr,
    )
    sys.exit(1)

REPO_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = REPO_ROOT / "docs" / "GOVGIG_ARCHITECTURE.md"
OUT_PATH = REPO_ROOT / "docs" / "GOVGIG_Architecture_And_Env.docx"


def read_md():
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Source not found: {MD_PATH}")
    return MD_PATH.read_text(encoding="utf-8")


def extract_mermaid_blocks(text: str) -> list[str]:
    blocks = []
    pattern = re.compile(r"```mermaid\s*\n(.*?)```", re.DOTALL)
    for m in pattern.finditer(text):
        blocks.append(m.group(1).strip())
    return blocks


def extract_narrative(text: str) -> str:
    start = "## 2. Architecture Narrative"
    end = "## 3. Environment Variables"
    i = text.find(start)
    j = text.find(end)
    if i == -1 or j == -1:
        return ""
    snippet = text[i:j]
    lines = snippet.split("\n")
    out = []
    for line in lines:
        if line.strip() == start.strip():
            continue
        out.append(line)
    return "\n".join(out).strip()


def extract_table_rows(text: str) -> list[list[str]]:
    in_table = False
    rows = []
    for line in text.split("\n"):
        if "## 3. Environment Variables" in line:
            in_table = True
            continue
        if not in_table:
            continue
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if not cells:
            continue
        # Skip separator line (Variable | Used by | Purpose)
        if all(
            set(c.replace("-", "").strip()) <= set(" ") or c.strip() == ""
            for c in cells
        ):
            continue
        rows.append(cells)
    return rows


def add_paragraph(doc, text: str, style=None):
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_heading(doc, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_code(doc, code: str, font_size: Pt = Pt(9)):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = font_size
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def build_docx(md_text: str) -> Document:
    doc = Document()
    mermaid_blocks = extract_mermaid_blocks(md_text)
    narrative = extract_narrative(md_text)
    table_rows = extract_table_rows(md_text)

    add_heading(doc, "GovGig AI – Architecture and Environment Reference", level=0)
    doc.add_paragraph(
        "Scope: Backend AI system and orchestration layer (FedRAMP-relevant boundary). "
        "Excludes frontend/UI and cloud deployment; infrastructure will be added when AWS architecture is finalized."
    )

    add_heading(doc, "1. Architecture Diagram", level=1)
    doc.add_paragraph(
        "The architecture is defined in Mermaid in the repo (docs/GOVGIG_ARCHITECTURE.md). "
        "To include it in this document: copy the Mermaid block from that file, paste at https://mermaid.live, "
        "export as PNG, and insert the image here."
    )
    if mermaid_blocks:
        add_heading(doc, "1.1 System Boundary (In-Scope Components)", level=2)
        add_code(doc, mermaid_blocks[0])
        if len(mermaid_blocks) > 1:
            add_heading(doc, "1.2 Request Data Flow", level=2)
            add_code(doc, mermaid_blocks[1])

    add_heading(doc, "2. Architecture Narrative", level=1)
    for block in narrative.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        add_paragraph(doc, block)

    add_heading(doc, "3. Environment Variables", level=1)
    doc.add_paragraph(
        "Source of truth: Backend src/config.py, Ingest ingest_python/config.py. "
        "Example values: .env.example, infra/terraform.tfvars.example."
    )
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=3)
        table.style = "Table Grid"
        for i, row_cells in enumerate(table_rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_cells[:3]):
                row.cells[j].text = cell_text
        doc.add_paragraph()

    return doc


def main():
    md_text = read_md()
    doc = build_docx(md_text)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
